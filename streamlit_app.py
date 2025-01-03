import streamlit as st
import pandas as pd
import numpy as np
from scipy.stats import ttest_ind
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
from docx import Document

# Function to create a Word document
def create_word_doc(content):
    doc = Document()
    for section in content:
        doc.add_heading(section['title'], level=1)
        for table in section.get('tables', []):
            doc.add_paragraph(f"Table: {table['title']}")
            df = table['dataframe']
            table_doc = doc.add_table(rows=1, cols=len(df.columns))
            table_doc.style = 'Table Grid'
            # Add headers
            hdr_cells = table_doc.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = str(col)
            # Add rows
            for _, row in df.iterrows():
                row_cells = table_doc.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        for chart in section.get('charts', []):
            doc.add_paragraph(f"Chart: {chart['title']}")
            chart['figure'].savefig(chart['filename'])
            doc.add_picture(chart['filename'])
    return doc

# Upload data
st.title("Streamlit Data Analysis App")
uploaded_file = st.file_uploader("Upload your Excel file", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    export_content = []

    # Tab structure
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["Distribution Tables", "Pivot Tables", "Statistical Analysis", "Correlations", "Graph Builder"])

    # Tab 1: Distribution Tables
    with tab1:
        st.header("Automated Distribution Tables")
        tab1_content = {"title": "Distribution Tables", "tables": [], "charts": []}

        for column in df.columns:
            if df[column].dtype in [np.int64, np.float64, object]:
                st.subheader(f"Distribution for {column}")
                use_ranges = st.checkbox(f"Use Ranges for {column}?", key=f"{column}_ranges")
                if use_ranges and df[column].dtype in [np.int64, np.float64]:
                    range_step = st.number_input(f"Step size for {column} ranges", min_value=1, value=10, key=f"{column}_range_step")
                    bins = list(range(int(df[column].min()), int(df[column].max()) + range_step, range_step))
                    if len(bins) > 1:
                        labels = [f"{bins[i]}-{bins[i+1]-1}" for i in range(len(bins)-1)]
                        df[column] = pd.cut(df[column], bins=bins, labels=labels, right=False)
                    else:
                        st.warning(f"Invalid range for column {column}. Skipping range binning.")
                        continue

                distribution = df[column].value_counts().reset_index()
                distribution.columns = [column, "Count"]
                distribution["Percentage"] = (distribution["Count"] / distribution["Count"].sum() * 100).round(2).astype(str) + '%'
                distribution.reset_index(drop=True, inplace=True)
                distribution.index = distribution.index + 1  # Start index from 1

                if not distribution.empty:
                    total_row = pd.DataFrame({column: ["Total"], "Count": [distribution["Count"].sum()], "Percentage": ["100%"]})
                    distribution = pd.concat([distribution, total_row], ignore_index=True)

                    st.dataframe(distribution)
                    tab1_content["tables"].append({"title": f"Distribution for {column}", "dataframe": distribution})

                    # Plot chart
                    if len(distribution) > 1:
                        fig, ax = plt.subplots()
                        distribution.iloc[:-1].plot(kind="bar", x=column, y="Count", ax=ax, legend=False)
                        ax.set_title(f"{column} Distribution")
                        st.pyplot(fig)
                    else:
                        st.info(f"No data available to plot for column {column}.")
                else:
                    st.info(f"No data available for column {column}.")

        export_content.append(tab1_content)

    # Tab 2: Pivot Tables
    with tab2:
        st.header("Pivot Tables")
        st.write("Select columns to create pivot tables.")
        rows = st.multiselect("Rows", df.columns)
        cols = st.multiselect("Columns", df.columns)
        values = st.selectbox("Values", df.columns)
        agg_func = st.selectbox("Aggregation Function", ["mean", "sum", "count", "max", "min"])
        filters = st.multiselect("Filters", df.columns)

        if st.button("Generate Pivot Table"):
            try:
                pivot_table = pd.pivot_table(df, index=rows, columns=cols, values=values, aggfunc=agg_func)
                pivot_table.reset_index(drop=True, inplace=True)
                pivot_table.index = pivot_table.index + 1  # Start index from 1
                st.dataframe(pivot_table)
            except Exception as e:
                st.error(f"Error generating pivot table: {e}")

    # Tab 3: Statistical Analysis
    with tab3:
        st.header("Statistical Analysis")
        st.write("Select columns for statistical calculations.")
        selected_columns = st.multiselect("Select Columns", df.select_dtypes(include=[np.number]).columns)
        if len(selected_columns) >= 2:
            col1, col2 = selected_columns[:2]
            st.write(f"Calculating statistics between {col1} and {col2}")

            # Compute statistics
            stats = {
                "Metric": ["Mean", "Median", "Std Dev", "T-Statistic", "P-Value"],
                col1: [
                    df[col1].mean(),
                    df[col1].median(),
                    df[col1].std(),
                    None,  # Placeholder for T-Statistic
                    None,  # Placeholder for P-Value
                ],
                col2: [
                    df[col2].mean(),
                    df[col2].median(),
                    df[col2].std(),
                    None,
                    None,
                ],
            }
            t_stat, p_value = ttest_ind(df[col1].dropna(), df[col2].dropna())
            stats[col1][3] = t_stat
            stats[col1][4] = p_value
            stats[col2][3] = t_stat
            stats[col2][4] = p_value

            stats_df = pd.DataFrame(stats)
            stats_df.index = stats_df.index + 1  # Start index from 1
            st.dataframe(stats_df)

    # Tab 4: Correlations
    with tab4:
        st.header("Correlations")
        st.write("Correlation matrix of numeric columns.")
        numeric_df = df.select_dtypes(include=[np.number])
        if numeric_df.empty:
            st.warning("No numeric columns available for correlation.")
        else:
            correlation_matrix = numeric_df.corr()
            correlation_matrix.reset_index(drop=True, inplace=True)
            correlation_matrix.index = correlation_matrix.index + 1  # Start index from 1
            st.dataframe(correlation_matrix)

            fig, ax = plt.subplots()
            sns.heatmap(correlation_matrix, annot=True, cmap="coolwarm", ax=ax)
            ax.set_title("Correlation Heatmap")
            st.pyplot(fig)

    # Tab 5: Graph Builder
    with tab5:
        st.header("Graph Builder")
        st.write("Select columns to build graphs.")
        x_col = st.selectbox("X-Axis", df.columns)
        y_col = st.selectbox("Y-Axis", df.columns)
        graph_type = st.selectbox("Graph Type", ["Scatter", "Line", "Bar", "Histogram", "Boxplot"])

        graph_title = st.text_input("Graph Title", value=f"{x_col} vs {y_col}")
        x_label = st.text_input("X-Axis Label", value=x_col)
        y_label = st.text_input("Y-Axis Label", value=y_col)

        if st.button("Generate Graph"):
            fig, ax = plt.subplots()
            if graph_type == "Scatter":
                sns.scatterplot(x=df[x_col], y=df[y_col], ax=ax)
            elif graph_type == "Line":
                sns.lineplot(x=df[x_col], y=df[y_col], ax=ax)
            elif graph_type == "Bar":
                sns.barplot(x=df[x_col], y=df[y_col], ax=ax)
            elif graph_type == "Histogram":
                sns.histplot(df[x_col], bins=30, kde=True, ax=ax)
            elif graph_type == "Boxplot":
                sns.boxplot(x=df[x_col], y=df[y_col], ax=ax)
            ax.set_title(graph_title)
            ax.set_xlabel(x_label)
            ax.set_ylabel(y_label)
            st.pyplot(fig)

    # Download Button
    if st.button("Download as Word Document"):
        doc = create_word_doc(export_content)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("Download Word Document", buffer, "data_analysis.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
